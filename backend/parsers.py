"""Resume parsing (PDF/DOCX) and document export (txt/pdf/docx)."""
import io
import pdfplumber
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch


def parse_pdf(data: bytes) -> str:
    out = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            out.append(text)
    return "\n".join(out).strip()


def parse_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    paras = [p.text for p in doc.paragraphs]
    return "\n".join(paras).strip()


def parse_resume_file(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return parse_pdf(data)
    if name.endswith(".docx"):
        return parse_docx(data)
    if name.endswith(".doc"):
        # Best-effort decode for .doc fallback
        return data.decode("utf-8", errors="ignore").strip()
    if name.endswith(".txt"):
        return data.decode("utf-8", errors="ignore").strip()
    # Fallback: try utf-8
    return data.decode("utf-8", errors="ignore").strip()


def export_txt(text: str) -> bytes:
    return text.encode("utf-8")


def export_docx(text: str) -> bytes:
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_pdf(text: str) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=0.7 * inch, rightMargin=0.7 * inch,
        topMargin=0.7 * inch, bottomMargin=0.7 * inch,
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "body", parent=styles["Normal"],
        fontName="Helvetica", fontSize=10.5, leading=14,
    )
    story = []
    for raw in text.split("\n"):
        line = raw.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if line.strip() == "":
            story.append(Spacer(1, 6))
        else:
            story.append(Paragraph(line, body))
    doc.build(story)
    return buf.getvalue()
